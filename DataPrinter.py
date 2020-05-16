import os

import plotly.graph_objects as go
import numpy
from docx.shared import Inches

threshold_KO_I = 1.96
threshold_KO_II = 1.96
header_color = 'lightgray'
regular = 'rgb(239, 243, 255)'
light = 'lightgreen'
extreme = 'rgb(200, 70, 50)'
medium = 'yellow'
scale = [[0, 0.3, regular], [0.3, 0.5, light], [0.5, 0.7, medium], [0.7, 1, extreme]]



# save KO_I statistic table to path
# return list of cells in which the threshold value is exceeded
def print_table_KO_I(data, header, path, title, threshold):
    threshold_exceeded = []
    formatted_data = list(numpy.copy(data))
    formatted_data.insert(0, header)
    formatted_header = list(numpy.copy(header))
    formatted_header.insert(0, '')

    # fill_color
    fill_color = []
    for i in range(len(formatted_data)):
        line = []
        for j in range(len(formatted_data[i])):
            if i == 0:
                line.append(header_color)
            else:
                if formatted_data[i][j] <= threshold:
                    line.append(regular)
                else:
                    line.append(extreme)
                    threshold_exceeded.append((header[i - 1], header[j]))
        fill_color.append(line)

    fig = go.Figure(data=[
        go.Table(header=dict(values=formatted_header,
                             fill_color=header_color,
                             ),
                 cells=dict(values=formatted_data,
                            fill_color=fill_color))
    ])
    fig.update_layout(
        title=title,
    )
    fig.write_image(path)
    return threshold_exceeded


def print_distractor_frequency_stat(data, path, title):
    formatted_data = []
    formatted_header = []
    for i in range(len(data)):
        formatted_data.append(data[i][0])
        formatted_header.append('answer ' + str(data[i][1]))
    fig = go.Figure([go.Bar(x=formatted_header, y=formatted_data)])
    fig.update_layout(
        title=title,
        yaxis=dict(
            title='count',
        )
    )
    fig.write_image(path)


def create_report_KO_I(document, formulation_homogeneity, path):
    if not os.path.exists(path):
        os.mkdir(path)

    for key, question_stat in formulation_homogeneity.items():
        pic_title = "КО-I statistic question {}".format(key)
        pic_path = os.path.join(path, "КО-I {}.jpg".format(key))
        threshold_exceeded = print_table_KO_I(question_stat[0],
                                              list(question_stat[1]),
                                              pic_path,
                                              pic_title,
                                              threshold_KO_I)
        document.add_picture(pic_path, width=Inches(5))
        for pair in threshold_exceeded:
            document.add_paragraph(
                'Варианты формулировок {} и {} неоднородны, их нужно исключить из рассмотрения.'
                    .format(pair[0], pair[1]))
    return document


def create_report_KO_II(document, distractor_frequency_stat, distractor_homogeneity, path):
    if not os.path.exists(path):
        os.mkdir(path)

    for question, question_stat in distractor_frequency_stat.items():
        for formulation, formulation_stat in question_stat.items():
            pic_title = "distractor_frequency_stat question {} - formulation {}".format(question, formulation)
            pic_path = os.path.join(path, "{}.jpg".format(pic_title))
            print_distractor_frequency_stat(formulation_stat, pic_path, pic_title)
            document.add_picture(pic_path, width=Inches(5))

        if distractor_homogeneity[question] is None:
            document.add_paragraph(
                'В вопросе {} проверка неоднородности распределения дистракторов невозможна. '
                'Следует изменить и переработать структуру дистракторов формулировок тестового задания.'.format(
                    question))
        else:
            if distractor_homogeneity[question][0] > distractor_homogeneity[question][1]:
                document.add_paragraph(
                    'В вопросе {} распределение дистракторов неоднородно. '
                    'Следует изменить и переработать структуру дистракторов формулировок тестового задания.'.format(
                        question))
            else:
                document.add_paragraph(
                    'В вопросе {} распределение дистракторов однородно.'.format(question))
    return document


def create_report_correlation(document, correlation_stat, path):
    if not os.path.exists(path):
        os.mkdir(path)

    for key, question_matrix in correlation_stat.items():
        pic_title = "question correlation - test {}".format(key)
        pic_path = os.path.join(path, "correlation {}.jpg".format(key))
        threshold_exceeded, threshold_medium = print_table_correlation(question_matrix[0],
                                                     list(question_matrix[1]),
                                                     pic_path,
                                                     pic_title)
        document.add_picture(pic_path, width=Inches(5))
        for pair in threshold_exceeded:
            document.add_paragraph(
                'Вопросы {} и {} сильно коррелируют, один из вопросов следует исключить из теста.'
                    .format(pair[0], pair[1]))
        for pair in threshold_medium:
            document.add_paragraph(
                'Вопросы {} и {} заметно коррелируют, их нельзя предъявлять друг за другом.'
                    .format(pair[0], pair[1]))
    return document


def print_table_correlation(data, header, path, title):
    threshold_exceeded = []
    threshold_medium = []
    formatted_data = list(numpy.copy(data))
    formatted_data.insert(0, header)
    formatted_header = list(numpy.copy(header))
    formatted_header.insert(0, '')

    # fill_color
    fill_color = []
    for i in range(len(formatted_data)):
        line = []
        for j in range(len(formatted_data[i])):
            if i == 0:
                line.append(header_color)
            else:
                if formatted_data[i][j] == "-" or abs(formatted_data[i][j]) == 1:
                    line.append(regular)
                else:
                    for s_index in range(len(scale)):
                        s = scale[s_index]
                        if s[0] <= abs(formatted_data[i][j]) < s[1]:
                            line.append(s[2])
                            if s_index == 3:
                                threshold_exceeded.append((header[j], header[i - 1]))
                            if s_index == 2:
                                threshold_medium.append((header[j], header[i - 1]))

        fill_color.append(line)

    fig = go.Figure(data=[
        go.Table(header=dict(values=formatted_header,
                             fill_color=header_color,
                             ),
                 cells=dict(values=formatted_data,
                            fill_color=fill_color))
    ])
    fig.update_layout(
        title=title,
    )
    fig.write_image(path)
    return threshold_exceeded, threshold_medium

